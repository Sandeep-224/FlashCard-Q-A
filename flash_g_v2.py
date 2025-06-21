import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext, simpledialog
import os
import datetime
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import pyperclip
import fitz  # PyMuPDF
from transformers import T5Tokenizer, T5ForConditionalGeneration

# Load model & tokenizer
tokenizer = T5Tokenizer.from_pretrained("valhalla/t5-base-qg-hl")
model = T5ForConditionalGeneration.from_pretrained("valhalla/t5-base-qg-hl")

# Themes
LIGHT_THEME = {
    "bg": "#f5f0e6", "fg": "#4b3e2a", "btn_bg": "#d6c3b4", "btn_active": "#c7b4a3",
    "text_bg": "#fffaf5", "text_fg": "#4b3e2a", "entry_bg": "#fffaf5", "entry_fg": "#4b3e2a", "border": "#b8a38c"
}

DARK_THEME = {
    "bg": "#2c241b", "fg": "#e0d4c3", "btn_bg": "#5a4b3c", "btn_active": "#7a6957",
    "text_bg": "#3d3429", "text_fg": "#e0d4c3", "entry_bg": "#3d3429", "entry_fg": "#e0d4c3", "border": "#7a6a5a"
}

class FlashcardApp:
    def __init__(self, root):
        self.root = root
        self.theme = LIGHT_THEME
        self.flashcards = []

        self.root.title("📘 AI-Powered Flashcard Generator")
        self.root.geometry("900x650")
        self.root.resizable(True, True)

        self.apply_theme()
        self.build_layout()

    def apply_theme(self):
        self.root.configure(bg=self.theme["bg"])

    def build_layout(self):
        style = ttk.Style()
        style.theme_use('clam')
        style.configure("TButton", font=("Segoe UI", 10), padding=6)

        self.title_label = ttk.Label(
            self.root, text="Generate Q&A Flashcards from Documents",
            font=("Segoe UI", 16, "bold"),
            background=self.theme["bg"], foreground=self.theme["fg"]
        )
        self.title_label.pack(pady=(10, 5))

        # Center-aligned control buttons
        self.control_frame = ttk.Frame(self.root)
        self.control_frame.pack(pady=10)

        for widget in self.control_frame.winfo_children():
            widget.destroy()

        btns = [
            ("📂 Choose File", self.load_file),
            ("🌓 Toggle Theme", self.toggle_theme),
            ("📋 Copy Output", self.copy_output),
            ("📤 Export Output", self.export_output)
        ]

        for text, command in btns:
            btn = ttk.Button(self.control_frame, text=text, command=command)
            btn.pack(side=tk.LEFT, padx=10)

        # Flashcard output display
        self.text_frame = ttk.Frame(self.root)
        self.text_frame.pack(expand=True, fill=tk.BOTH, padx=10, pady=5)

        self.output_box = scrolledtext.ScrolledText(
            self.text_frame, wrap=tk.WORD, font=("Consolas", 11),
            bg=self.theme["text_bg"], fg=self.theme["text_fg"],
            insertbackground=self.theme["text_fg"], relief="solid", borderwidth=2
        )
        self.output_box.pack(expand=True, fill=tk.BOTH)

        # Status bar
        self.status_var = tk.StringVar()
        self.status_bar = ttk.Label(
            self.root, textvariable=self.status_var, anchor="w", relief="sunken",
            background=self.theme["bg"], foreground=self.theme["fg"], font=("Segoe UI", 9)
        )
        self.status_bar.pack(fill=tk.X, side=tk.BOTTOM)
        self.set_status("Ready")

    def set_status(self, msg):
        self.status_var.set(msg)
        self.root.update_idletasks()

    def toggle_theme(self):
        self.theme = DARK_THEME if self.theme == LIGHT_THEME else LIGHT_THEME
        self.apply_theme()
        self.update_widget_colors()

    def update_widget_colors(self):
        self.root.configure(bg=self.theme["bg"])
        self.title_label.configure(background=self.theme["bg"], foreground=self.theme["fg"])
        self.output_box.configure(bg=self.theme["text_bg"], fg=self.theme["text_fg"], insertbackground=self.theme["text_fg"])
        self.status_bar.configure(background=self.theme["bg"], foreground=self.theme["fg"])

    def load_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Text files", "*.txt"), ("Word files", "*.docx"), ("PDF files", "*.pdf")])
        if file_path:
            try:
                self.set_status("Processing file...")
                text = extract_text(file_path)
                self.output_box.delete(1.0, tk.END)
                self.output_box.insert(tk.END, "Generating flashcards...\n")
                self.root.update_idletasks()

                self.flashcards = generate_flashcards(text)
                self.display_flashcards(self.flashcards)
                self.set_status("Flashcards generated successfully.")
            except Exception as e:
                self.set_status("Error occurred.")
                messagebox.showerror("Error", f"Failed to process file:\n{e}")

    def display_flashcards(self, flashcards):
        self.output_box.delete(1.0, tk.END)
        if not flashcards:
            self.output_box.insert(tk.END, "No flashcards could be generated.")
            return
        for i, (q, a) in enumerate(flashcards, 1):
            self.output_box.insert(tk.END, f"Flashcard #{i}\nQ: {q}\nA: {a}\n\n")

    def copy_output(self):
        text = self.output_box.get("1.0", tk.END).strip()
        if text:
            pyperclip.copy(text)
            self.set_status("Output copied to clipboard.")
        else:
            self.set_status("Nothing to copy.")

    def export_output(self):
        if not self.flashcards:
            messagebox.showinfo("No data", "Please generate flashcards first.")
            return

        format_choice = simpledialog.askstring("Export Format", "Export as PDF or Word? (Type 'pdf' or 'word')").strip().lower()
        if format_choice not in {"pdf", "word"}:
            messagebox.showerror("Invalid format", "Please enter 'pdf' or 'word'.")
            return

        filetypes = [("PDF files", "*.pdf")] if format_choice == "pdf" else [("Word Document", "*.docx")]
        ext = ".pdf" if format_choice == "pdf" else ".docx"

        file_path = filedialog.asksaveasfilename(defaultextension=ext, filetypes=filetypes)
        if not file_path:
            return

        if format_choice == "pdf":
            export_to_pdf(file_path, self.flashcards)
        else:
            export_to_word(file_path, self.flashcards)

        self.set_status(f"Flashcards exported as {format_choice.upper()}.")

# --- Utility Functions ---

def extract_text_from_docx(path):
    doc = DocxDocument(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

def extract_text_from_pdf(path):
    text = ""
    with fitz.open(path) as doc:
        for page in doc:
            text += page.get_text()
    return text

def extract_text(path):
    if path.endswith(".txt"):
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    elif path.endswith(".docx"):
        return extract_text_from_docx(path)
    elif path.endswith(".pdf"):
        return extract_text_from_pdf(path)
    else:
        raise ValueError("Unsupported file type.")

def get_relevant_paragraphs(text, min_len=80, max_len=400, limit=25):
    return [p.strip() for p in text.split("\n") if min_len <= len(p.strip()) <= max_len][:limit]

def generate_questions(text):
    input_text = "generate questions: " + text
    encoding = tokenizer.encode_plus(input_text, return_tensors="pt", padding=True, truncation=True, max_length=512)
    outputs = model.generate(**encoding, max_length=64, num_beams=4, early_stopping=True)
    return [tokenizer.decode(out, skip_special_tokens=True) for out in outputs]

def generate_flashcards(text, max_flashcards=25):
    flashcards = []
    for para in get_relevant_paragraphs(text, limit=max_flashcards):
        try:
            for q in generate_questions(para):
                flashcards.append((q, para))
        except Exception as e:
            print(f"Error generating for paragraph: {e}")
    return flashcards

def export_to_word(path, flashcards):
    doc = DocxDocument()
    doc.add_heading("Flashcards", level=1)
    doc.add_paragraph(f"Generated on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}\n")

    for i, (q, a) in enumerate(flashcards, 1):
        doc.add_paragraph(f"Flashcard #{i}", style="Heading2")
        doc.add_paragraph(f"Q: {q}")
        doc.add_paragraph(f"A: {a}")
        doc.add_paragraph("")

    doc.save(path)

def export_to_pdf(path, flashcards):
    c = canvas.Canvas(path, pagesize=letter)
    width, height = letter
    y = height - 50
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y, "Flashcards")
    c.setFont("Helvetica", 10)
    y -= 25
    c.drawString(50, y, f"Generated on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
    y -= 30

    for i, (q, a) in enumerate(flashcards, 1):
        if y < 100:
            c.showPage()
            y = height - 50
        c.setFont("Helvetica-Bold", 11)
        c.drawString(50, y, f"Flashcard #{i}")
        y -= 15
        c.setFont("Helvetica", 10)
        c.drawString(50, y, f"Q: {q}")
        y -= 15
        c.drawString(50, y, f"A: {a}")
        y -= 25

    c.save()

if __name__ == "__main__":
    root = tk.Tk()
    app = FlashcardApp(root)
    root.mainloop()
